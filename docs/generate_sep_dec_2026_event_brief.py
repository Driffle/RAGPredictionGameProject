#!/usr/bin/env python3
"""Export Floor Brief top-event desks to PDF and DOCX.

Layout mirrors the website desk: dark mast, kicker, event lede cards,
runtime KPIs, social engagement tags, and a mapped-products grid of
thirty recommended games per event. Pass --start/--end/--stem to build
another window; defaults keep the Sep–Dec 2026 brief.
"""

from __future__ import annotations

import argparse
import html
import io
import re
import shutil
import sys
from datetime import date
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.calendar_dedupe import canonical_event_name, dedupe_calendar_rows, is_gaming_world_event, unique_event_listings
from src.date_range import event_window, events_in_range, row_precision
from src.dates import annotate_event
from src.content_marketing import FAMILY_TAGS, PLATFORMS, UNIVERSE_TAGS, correlation_from_plan, hashtag, social_pack
from src.first_party import showcase_owner
from src.load_data import load_catalog
from src.match import build_title_index, superhero_universe_for_row
from src.promote import calendar_label, promo_family, recommended_games_for_event

RANGE_START = date(2026, 9, 1)
RANGE_END = date(2026, 12, 31)
STEM = "FloorBrief_Sep-Dec_2026_Top_Events"
DESKTOP = Path.home() / "Desktop"
DOCS = ROOT / "docs"
GAMES_PER_EVENT = 30

BG = "07080F"
INK = "F4F7FF"
MUTED = "9AA6C3"
CYAN = "5CE1E6"
MAGENTA = "FF5CA8"
AMBER = "FFC14A"
CARD = "0E1222"
LINE = "2A3148"

MAJOR_WEIGHTS = (
    ("the game awards", 90),
    ("gamescom", 87),
    ("summer game fest", 86),
    ("e3", 85),
    ("state of play", 84),
    ("nintendo direct", 83),
    ("tokyo game show", 80),
    ("steam next fest", 78),
    ("xbox games showcase", 76),
    ("pax west", 50),
    ("blizzcon", 60),
    ("paris games week", 58),
    ("brasil game show", 52),
    ("golden joystick", 48),
    ("g-star", 46),
)
E3_NAME = re.compile(r"\be3\b", re.I)


def nice_date(value: str | None) -> str:
    text = (value or "")[:10]
    if len(text) < 10:
        return value or "—"
    parsed = date.fromisoformat(text)
    return parsed.strftime("%-d %b %Y")


def runtime_label(row: dict) -> str:
    dated = annotate_event(row)
    label = dated.get("date_label") or ""
    if label:
        return label
    start, end = event_window(row)
    if not start:
        return "—"
    left = nice_date(start.isoformat())
    right = nice_date((end or start).isoformat())
    return left if left == right else f"{left} → {right}"


def load_calendar() -> list[dict]:
    from src.database import live_events

    rows = live_events()
    if rows:
        return rows
    from src.historical_calendar import historical_events
    from src.horizon import projected_events
    from src.load_data import load_events

    return list(historical_events()) + list(projected_events()) + list(load_events())


def _name_matches(name: str, needle: str) -> bool:
    if needle == "e3":
        return bool(E3_NAME.search(name))
    return needle in name


def event_family(row: dict) -> str:
    name = (row.get("event") or "").lower()
    if "the game awards" in name:
        return "the game awards"
    if "gamescom" in name:
        return "gamescom"
    if "summer game fest" in name:
        return "summer game fest"
    if "state of play" in name:
        return "state of play"
    if "nintendo direct" in name:
        return "nintendo direct"
    if E3_NAME.search(name):
        return "e3"
    if "steam next fest" in name:
        return "steam next fest"
    if "tokyo game show" in name:
        return "tokyo game show"
    if "xbox games showcase" in name or "xbox developer direct" in name:
        return "xbox showcase"
    return canonical_event_name(row)


def event_score(row: dict) -> tuple:
    name = (row.get("event") or "").lower()
    event_type = (row.get("event_type") or "").lower()
    status = f"{row.get('confirmation') or ''} {row.get('status') or ''}".lower()
    score = 0
    for needle, points in MAJOR_WEIGHTS:
        if _name_matches(name, needle):
            score += points
            break
    else:
        if "showcase" in event_type or "awards" in event_type:
            score += 30
        elif "expo" in event_type:
            score += 24
        elif "festival" in event_type:
            score += 18
    if "japan" in name and "state of play" in name:
        score -= 20
    if "partner" in name:
        score -= 40
    if "concert" in name or "hollywood bowl" in name:
        score -= 60
    if "opening night" in name or "future games show" in name:
        score -= 16
    if "confirm" in status:
        score += 18
    precision = row_precision(row)
    if precision == "day":
        score += 8
    elif precision == "month":
        score -= 6
    start = row.get("start_date") or "0000-01-01"
    return (-score, tuple(-ord(ch) for ch in start), name)


def top_events(rows: list[dict], *, limit: int = 5) -> list[dict]:
    matched = events_in_range(rows, RANGE_START, RANGE_END, kind="event", precision="dated")
    gaming = [row for row in matched if is_gaming_world_event(row)]
    gaming = unique_event_listings(dedupe_calendar_rows(gaming))
    gaming.sort(key=event_score)
    picked: list[dict] = []
    seen_family: set[str] = set()
    for row in gaming:
        family = event_family(row)
        if not family or family in seen_family:
            continue
        seen_family.add(family)
        picked.append(row)
        if len(picked) >= limit:
            break
    picked.sort(key=lambda row: row.get("start_date") or "9999")
    return picked


def _unique_tags(*groups: list) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for group in groups:
        for tag in group or []:
            if tag and tag not in seen:
                seen.add(tag)
                out.append(tag)
    return out


def event_engagement(row: dict, games: list[dict]) -> dict:
    name = calendar_label(row)
    family = promo_family(row)
    universe = superhero_universe_for_row(name)
    hero = (games[0].get("canonical_title") if games else "") or name
    pack = social_pack(hero, name, family)
    tags = _unique_tags(
        [hashtag(name), "#ad"],
        UNIVERSE_TAGS.get(universe) or [],
        FAMILY_TAGS.get(family) or FAMILY_TAGS["default"],
        (pack.get("tiktok") or {}).get("hashtags") or [],
    )
    platforms = {}
    for key, label, when in PLATFORMS:
        item = pack.get(key) or {}
        platforms[key] = {
            "label": item.get("platform") or label,
            "hashtags": (item.get("hashtags") or [])[:8],
            "best_times": item.get("best_times") or when,
        }
    return {
        "family": family,
        "hashtags": tags[:10],
        "platforms": platforms,
    }


def product_engagement(game: dict, event_name: str, family: str, start: str, end: str) -> dict:
    corr = (
        correlation_from_plan(
            {
                "canonical_title": game.get("canonical_title") or "",
                "event": event_name,
                "platform": game.get("platform") or "",
                "promo_family": family,
                "role": "game",
                "runtime_start": start,
                "runtime_end": end,
                "promo_start": start,
                "promo_end": end,
            }
        )
        or {}
    )
    social = corr.get("social") or {}
    return {
        "hashtags": (corr.get("top_hashtags") or [])[:5],
        "seo_keywords": (corr.get("seo_keywords") or [])[:4],
        "post_times": corr.get("post_times") or {},
        "platforms": {
            key: {
                "label": pack.get("platform") or key,
                "hashtags": (pack.get("hashtags") or [])[:6],
                "best_times": pack.get("best_times") or "",
            }
            for key, pack in social.items()
        },
    }


def chip_html(tags: list[str]) -> str:
    if not tags:
        return '<span class="meta">—</span>'
    return "".join(f'<span class="hashtag">{html.escape(tag)}</span>' for tag in tags)


def build_payload() -> dict:
    catalog = load_catalog(games_only=True, drop_placeholder_dates=True)
    title_index = build_title_index(catalog)
    events = []
    for row in top_events(load_calendar()):
        games = recommended_games_for_event(row, catalog, limit=GAMES_PER_EVENT, title_index=title_index)
        start, end = event_window(row)
        owner = showcase_owner(calendar_label(row))
        start_iso = start.isoformat() if start else ""
        end_iso = (end or start).isoformat() if (end or start) else ""
        social = event_engagement(row, games)
        events.append(
            {
                "name": calendar_label(row),
                "kind": row.get("event_type") or row.get("kind") or "event",
                "status": row.get("confirmation") or row.get("status") or "planning",
                "location": row.get("location") or "",
                "country": row.get("country") or "",
                "mode": row.get("attendance_mode") or "",
                "related": row.get("related_game") or "",
                "runtime": runtime_label(row),
                "start": start_iso,
                "end": end_iso,
                "owner": owner.publisher_label if owner else "",
                "hashtags": social["hashtags"],
                "platforms": social["platforms"],
                "family": social["family"],
                "games": [
                    {
                        "title": game.get("canonical_title") or "",
                        "platform": game.get("platform") or "",
                        "release": nice_date(game.get("release_date") or "")
                        if game.get("release_date")
                        else "TBA",
                        "type": game.get("product_type") or "game",
                        "publisher": game.get("publisher") or "",
                        **product_engagement(game, calendar_label(row), social["family"], start_iso, end_iso),
                    }
                    for game in games
                ],
            }
        )
    start_label = nice_date(RANGE_START.isoformat())
    end_label = nice_date(RANGE_END.isoformat())
    return {
        "as_of": date.today().isoformat(),
        "range_label": f"{start_label} → {end_label}",
        "kicker": f"Live desk · {RANGE_START.year}–{RANGE_END.year}",
        "intro_title": "Which event windows are live?",
        "intro_body": (
            f"Top five merchandising events overlapping {start_label}–{end_label}, ranked like the Floor Brief desk: "
            "confirmed showcases and expos first. Each card lists the event runtime, social engagement tags "
            "(TikTok, Instagram, YouTube Shorts, X), and thirty catalog games to recommend — with product hashtags, "
            "post windows, and SEO keywords. Console showcases keep first-party owned-studio games "
            "(SIE on State of Play, Nintendo on Directs)."
        ),
        "intro_short": (
            f"Top five merchandising events overlapping {start_label}–{end_label}. "
            "Each card lists the event runtime, social engagement tags, and thirty catalog games to recommend. "
            "Console showcases keep first-party owned-studio games."
        ),
        "events": events,
    }


def render_html(payload: dict) -> str:
    cards = []
    for index, event in enumerate(payload["events"], start=1):
        tiles = []
        for rank, game in enumerate(event["games"], start=1):
            tiles.append(
                f"""
            <article class="media-tile">
              <span class="rank">#{rank:02d}</span>
              <h5>{html.escape(game["title"])}</h5>
              <p class="meta">{html.escape(game["release"])} · {html.escape(game["type"])}</p>
              <p class="meta">{html.escape(game["platform"] or "Multi")}</p>
              <div class="chip-row">{chip_html(game.get("hashtags") or [])}</div>
              <p class="meta seo">{html.escape(" · ".join(game.get("seo_keywords") or []) or "—")}</p>
            </article>"""
            )
        owner_badge = (
            f'<span class="badge on">{html.escape(event["owner"])}</span>' if event["owner"] else ""
        )
        platform_cols = "".join(
            f"""<div>
              <b>{html.escape(pack["label"])}</b>
              <p class="meta">Post when {html.escape(pack["best_times"])}</p>
              <div class="chip-row">{chip_html(pack.get("hashtags") or [])}</div>
            </div>"""
            for pack in (event.get("platforms") or {}).values()
        )
        cards.append(
            f"""
        <article class="lede">
          <div class="lede-hero">
            <div>
              <div class="badge-row">
                <span class="badge on">#{index:02d} · {html.escape(event["kind"])}</span>
                <span class="badge live">{html.escape(event["status"])}</span>
                {owner_badge}
                <span class="badge">{html.escape(event["mode"] or "digital / physical")}</span>
              </div>
              <h3 class="event-title">{html.escape(event["name"])}</h3>
              <p class="action">Event runtime {html.escape(event["runtime"])}</p>
              <p class="meta">{html.escape(event["location"] or event["country"] or "Worldwide")}</p>
              <div class="chip-row event-tags">{chip_html(event.get("hashtags") or [])}</div>
            </div>
          </div>
          <div class="kpis">
            <div class="kpi"><b>{html.escape(nice_date(event["start"]))}</b><span>Runtime start</span></div>
            <div class="kpi"><b>{html.escape(nice_date(event["end"]))}</b><span>Runtime end</span></div>
            <div class="kpi"><b>{len(event["games"])}</b><span>Recommended games</span></div>
            <div class="kpi"><b>{html.escape(event["mode"] or "—")}</b><span>Attendance</span></div>
          </div>
          <section class="card">
            <h4>Content marketing · social engagement</h4>
            <div class="content-social">{platform_cols}</div>
          </section>
          <section class="card">
            <h4>Mapped products · top {GAMES_PER_EVENT} recommended games</h4>
            <div class="media-grid">{"".join(tiles)}</div>
          </section>
        </article>"""
        )
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <title>Floor Brief · Sep–Dec 2026</title>
  <style>
    :root {{ --bg:#{BG}; --ink:#{INK}; --muted:#{MUTED}; --cyan:#{CYAN}; --magenta:#{MAGENTA}; --amber:#{AMBER}; --card:#{CARD}; --line:#{LINE}; }}
    * {{ box-sizing:border-box; }}
    html, body {{ margin:0; background:var(--bg); color:var(--ink); font:15px/1.5 "Helvetica Neue", Helvetica, Arial, sans-serif; }}
    body {{ -webkit-print-color-adjust:exact; print-color-adjust:exact; }}
    .aurora {{ position:fixed; inset:0; pointer-events:none; background:
      radial-gradient(900px 500px at 12% -10%, rgba(92,225,230,.22), transparent 55%),
      radial-gradient(700px 420px at 90% 8%, rgba(255,92,168,.18), transparent 50%),
      radial-gradient(600px 400px at 70% 90%, rgba(255,193,74,.1), transparent 55%); }}
    .mast, main {{ position:relative; z-index:1; }}
    .mast {{ display:flex; justify-content:space-between; align-items:flex-end; padding:28px 36px 18px; border-bottom:1px solid rgba(255,255,255,.12); }}
    .kicker {{ margin:0; letter-spacing:.16em; text-transform:uppercase; font-size:11px; color:var(--cyan); display:flex; align-items:center; gap:8px; }}
    .pulse-dot {{ width:8px; height:8px; border-radius:50%; background:var(--cyan); }}
    .mast h1 {{ margin:6px 0 0; font-family:Palatino, "Palatino Linotype", Georgia, serif; font-size:48px; letter-spacing:-.04em; line-height:.9;
      background:linear-gradient(120deg,#fff 20%,var(--cyan) 50%,var(--magenta) 90%); -webkit-background-clip:text; background-clip:text; color:transparent; }}
    .range {{ color:var(--muted); font-size:13px; text-align:right; }}
    main {{ padding:28px 36px 48px; width:min(1100px,100%); }}
    .intro h2 {{ font-family:Palatino, "Palatino Linotype", Georgia, serif; font-size:32px; margin:0 0 8px; letter-spacing:-.03em; }}
    .intro p {{ margin:0 0 22px; color:var(--muted); max-width:70ch; }}
    .lede, .card {{ background:rgba(14,18,34,.86); border:1px solid rgba(255,255,255,.12); border-radius:20px; padding:22px; margin:0 0 22px; }}
    .lede-hero {{ display:block; }}
    .event-title {{ margin:0 0 8px; font-family:Palatino, "Palatino Linotype", Georgia, serif; font-size:30px; letter-spacing:-.03em; }}
    .action {{ font-size:17px; color:var(--amber); margin:0 0 6px; }}
    .badge-row {{ display:flex; flex-wrap:wrap; gap:8px; margin:0 0 12px; }}
    .badge {{ font-size:10px; letter-spacing:.08em; text-transform:uppercase; padding:5px 9px; border-radius:999px; border:1px solid rgba(255,255,255,.12); }}
    .badge.on {{ border-color:var(--cyan); color:var(--cyan); }}
    .badge.live {{ border-color:#7dffb0; color:#7dffb0; }}
    .kpis {{ display:grid; grid-template-columns:repeat(4,1fr); gap:10px; margin:16px 0; }}
    .kpi {{ border:1px solid rgba(255,255,255,.12); border-radius:14px; padding:10px 12px; background:rgba(255,255,255,.03); }}
    .kpi b {{ display:block; font-size:16px; }}
    .kpi span {{ color:var(--muted); font-size:10px; letter-spacing:.1em; text-transform:uppercase; }}
    h4 {{ margin:0 0 12px; letter-spacing:.08em; text-transform:uppercase; font-size:11px; color:var(--cyan); }}
    .media-grid {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:10px; }}
    .media-tile {{ position:relative; padding:12px 14px; min-height:auto; border:1px solid rgba(255,255,255,.12); border-radius:16px; background:linear-gradient(145deg,rgba(255,255,255,.05),rgba(92,225,230,.04)); }}
    .media-tile h5 {{ margin:0 28px 4px 0; font:700 15px/1.25 Palatino, "Palatino Linotype", Georgia, serif; }}
    .meta {{ margin:3px 0; color:var(--muted); font-size:12px; }}
    .meta.seo {{ font-size:11px; }}
    .rank {{ position:absolute; right:10px; top:8px; color:var(--cyan); font-size:10px; letter-spacing:.12em; }}
    .chip-row {{ display:flex; flex-wrap:wrap; gap:4px; margin:6px 0 0; }}
    .hashtag {{ display:inline-block; font-size:11px; color:var(--cyan); background:rgba(92,225,230,.08); border:1px solid rgba(255,255,255,.12); border-radius:999px; padding:2px 8px; }}
    .content-social {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:10px; }}
    .content-social b {{ display:block; font-size:11px; letter-spacing:.08em; text-transform:uppercase; color:var(--cyan); margin-bottom:4px; }}
    .colophon {{ color:var(--muted); font-size:11px; padding:0 36px 28px; }}
    @page {{ size:A4; margin:12mm; }}
    @media print {{
      .lede {{ break-inside:auto; }}
      .media-tile, .card {{ break-inside:avoid; page-break-inside:avoid; }}
    }}
  </style>
</head>
<body>
  <div class="aurora"></div>
  <header class="mast">
    <div>
      <p class="kicker"><span class="pulse-dot"></span> {html.escape(payload["kicker"])}</p>
      <h1>Floor Brief</h1>
    </div>
    <p class="range">Event window<br><b style="color:{INK}">{html.escape(payload["range_label"])}</b><br>As of {html.escape(payload["as_of"])}</p>
  </header>
  <main>
    <div class="intro">
      <h2>{html.escape(payload["intro_title"])}</h2>
      <p>{html.escape(payload["intro_body"])}</p>
    </div>
    {"".join(cards)}
  </main>
  <p class="colophon">Floor Brief · mapped products from the live catalog and promotion plans · {html.escape(payload["as_of"])}</p>
</body>
</html>
"""


def write_pdf(payload: dict, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    bg = colors.HexColor(f"#{BG}")
    ink = colors.HexColor(f"#{INK}")
    muted = colors.HexColor(f"#{MUTED}")
    cyan = colors.HexColor(f"#{CYAN}")
    amber = colors.HexColor(f"#{AMBER}")
    card = colors.HexColor(f"#{CARD}")
    header = colors.HexColor("#12182A")
    stripe = colors.HexColor("#141A2C")

    kicker = ParagraphStyle("kicker", fontName="Helvetica-Bold", fontSize=8, textColor=cyan, leading=11, spaceAfter=2)
    title = ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=22, textColor=cyan, leading=26, spaceAfter=2)
    h2 = ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=14, textColor=ink, leading=18, spaceAfter=4)
    body = ParagraphStyle("body", fontName="Helvetica", fontSize=9, textColor=muted, leading=12)
    event_name = ParagraphStyle("event_name", fontName="Helvetica-Bold", fontSize=16, textColor=ink, leading=20, spaceAfter=2)
    runtime = ParagraphStyle("runtime", fontName="Helvetica-Bold", fontSize=11, textColor=amber, leading=14, spaceAfter=2)
    meta = ParagraphStyle("meta", fontName="Helvetica", fontSize=8, textColor=muted, leading=11)
    tags = ParagraphStyle("tags", fontName="Helvetica-Bold", fontSize=8, textColor=cyan, leading=11)
    cell = ParagraphStyle("cell", fontName="Helvetica", fontSize=8, textColor=ink, leading=10)
    cell_title = ParagraphStyle("cell_title", fontName="Helvetica-Bold", fontSize=8, textColor=ink, leading=10)
    cell_social = ParagraphStyle("cell_social", fontName="Helvetica", fontSize=7, textColor=cyan, leading=9)
    head_cell = ParagraphStyle("head_cell", fontName="Helvetica-Bold", fontSize=7, textColor=cyan, leading=9)
    range_style = ParagraphStyle("range", fontName="Helvetica-Bold", fontSize=10, textColor=ink, leading=13, alignment=TA_RIGHT)
    range_muted = ParagraphStyle("range_muted", fontName="Helvetica", fontSize=8, textColor=muted, leading=11, alignment=TA_RIGHT)
    left_wrap = ParagraphStyle("left_wrap", parent=kicker, alignment=TA_LEFT)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        title="Floor Brief · Sep–Dec 2026",
    )
    story = []

    mast = Table(
        [
            [
                [
                    Paragraph(payload["kicker"].upper(), left_wrap),
                    Paragraph("Floor Brief", title),
                ],
                [
                    Paragraph("Event window", range_muted),
                    Paragraph(payload["range_label"], range_style),
                    Paragraph(f"As of {payload['as_of']}", range_muted),
                ],
            ]
        ],
        colWidths=[110 * mm, 70 * mm],
    )
    mast.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), bg),
                ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(mast)
    story.append(Spacer(1, 8))
    story.append(Paragraph(payload["intro_title"], h2))
    story.append(Paragraph(html.escape(payload["intro_short"]), body))
    story.append(Spacer(1, 8))

    def text(value: str, style: ParagraphStyle) -> Paragraph:
        return Paragraph(html.escape(value or ""), style)

    for index, event in enumerate(payload["events"], start=1):
        owner = f"    ·    {event['owner']}" if event.get("owner") else ""
        place = event["location"] or event["country"] or "Worldwide"
        blocks = [
            text(f"#{index:02d}  {event['kind'].upper()}   ·   {event['status']}", kicker),
            text(event["name"], event_name),
            text(f"Event runtime  {event['runtime']}", runtime),
            text(
                f"{event['start']} → {event['end']}    ·    {place}    ·    {event['mode'] or '—'}{owner}",
                meta,
            ),
            text("  ".join(event.get("hashtags") or []) or "—", tags),
        ]
        for pack in (event.get("platforms") or {}).values():
            blocks.append(
                text(
                    f"{pack['label']}  ·  {pack['best_times']}  ·  {'  '.join(pack.get('hashtags') or [])}",
                    meta,
                )
            )
        header_table = Table([[blocks]], colWidths=[180 * mm])
        header_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), card),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )

        rows = [
            [
                text("#", head_cell),
                text("RECOMMENDED GAME", head_cell),
                text("RELEASE", head_cell),
                text("PLATFORM", head_cell),
                text("SOCIAL ENGAGEMENT", head_cell),
            ]
        ]
        for row_i, game in enumerate(event["games"], start=1):
            social_bits = "  ".join(game.get("hashtags") or [])
            seo = " · ".join(game.get("seo_keywords") or [])
            times = game.get("post_times") or {}
            bits = [
                times.get("tiktok") and f"TikTok · {times['tiktok']}",
                times.get("instagram") and f"IG · {times['instagram']}",
            ]
            extra = "  |  ".join(bit for bit in bits if bit)
            social = "<br/>".join(html.escape(part) for part in (social_bits, seo, extra) if part)
            rows.append(
                [
                    text(f"{row_i:02d}", cell),
                    text(game["title"], cell_title),
                    text(game["release"], cell),
                    text(game["platform"] or "Multi", cell),
                    Paragraph(social or "—", cell_social),
                ]
            )
        games_table = Table(rows, colWidths=[10 * mm, 58 * mm, 28 * mm, 32 * mm, 52 * mm], repeatRows=1)
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), header),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        for row_i in range(1, len(rows)):
            style_cmds.append(("BACKGROUND", (0, row_i), (-1, row_i), stripe if row_i % 2 else card))
        games_table.setStyle(TableStyle(style_cmds))
        story.append(header_table)
        story.append(Spacer(1, 3))
        story.append(games_table)
        story.append(Spacer(1, 10))

    story.append(
        Paragraph(
            f"Floor Brief · mapped products from the live catalog · {payload['as_of']}",
            meta,
        )
    )

    def paint(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(bg)
        canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
        canvas.restoreState()

    doc.build(story, onFirstPage=paint, onLaterPages=paint)


def _shade(cell, fill: str) -> None:
    from docx.oxml import parse_xml

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{fill}" w:val="clear"/>'
        )
    )


def _run(paragraph, text: str, *, size: int = 11, color: str = INK, bold: bool = False, font: str = "Calibri") -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.color.rgb = RGBColor.from_string(color)


def write_docx(payload: dict, path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    background = parse_xml(
        f'<w:background xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:color="{BG}"/>'
    )
    doc.element.insert(0, background)
    settings = doc.settings.element
    display = parse_xml(
        '<w:displayBackgroundShape xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    settings.append(display)

    def dark_table(rows: int, cols: int):
        table = doc.add_table(rows=rows, cols=cols)
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                _shade(cell, CARD)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
        return table

    mast = dark_table(1, 2)
    _shade(mast.cell(0, 0), BG)
    _shade(mast.cell(0, 1), BG)
    left = mast.cell(0, 0).paragraphs[0]
    _run(left, payload["kicker"].upper(), size=9, color=CYAN, bold=True)
    title = mast.cell(0, 0).add_paragraph()
    _run(title, "Floor Brief", size=28, color=CYAN, bold=True, font="Arial")
    right = mast.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(right, "Event window", size=9, color=MUTED)
    rng = mast.cell(0, 1).add_paragraph()
    rng.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rng, payload["range_label"], size=12, color=INK, bold=True)
    asof = mast.cell(0, 1).add_paragraph()
    asof.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(asof, f"As of {payload['as_of']}", size=10, color=MUTED)

    intro = dark_table(1, 1)
    _run(intro.cell(0, 0).paragraphs[0], payload["intro_title"], size=18, color=INK, bold=True, font="Arial")
    p = intro.cell(0, 0).add_paragraph()
    _run(
        p,
        payload["intro_short"],
        size=11,
        color=MUTED,
    )

    for index, event in enumerate(payload["events"], start=1):
        card = dark_table(1, 1)
        head = card.cell(0, 0).paragraphs[0]
        _run(head, f"#{index:02d}  {event['kind'].upper()}   ·   {event['status']}", size=9, color=CYAN, bold=True)
        name = card.cell(0, 0).add_paragraph()
        _run(name, event["name"], size=20, color=INK, bold=True, font="Arial")
        runtime = card.cell(0, 0).add_paragraph()
        _run(runtime, f"Event runtime  {event['runtime']}", size=13, color=AMBER, bold=True)
        meta = card.cell(0, 0).add_paragraph()
        _run(
            meta,
            f"{event['start']} → {event['end']}    ·    {event['location'] or event['country'] or 'Worldwide'}    ·    {event['mode'] or '—'}"
            + (f"    ·    {event['owner']}" if event["owner"] else ""),
            size=10,
            color=MUTED,
        )
        tags = card.cell(0, 0).add_paragraph()
        _run(tags, "  ".join(event.get("hashtags") or []) or "—", size=10, color=CYAN, bold=True)
        for pack in (event.get("platforms") or {}).values():
            line = card.cell(0, 0).add_paragraph()
            _run(
                line,
                f"{pack['label']}  ·  {pack['best_times']}  ·  {'  '.join(pack.get('hashtags') or [])}",
                size=9,
                color=MUTED,
            )

        games = dark_table(1 + len(event["games"]), 5)
        headers = ("#", "Recommended game", "Release", "Platform", "Social engagement")
        for col, label in enumerate(headers):
            _shade(games.cell(0, col), "12182A")
            para = games.cell(0, col).paragraphs[0]
            para.text = ""
            _run(para, label.upper(), size=8, color=CYAN, bold=True)
        for row_i, game in enumerate(event["games"], start=1):
            social_bits = "  ".join(game.get("hashtags") or [])
            seo = " · ".join(game.get("seo_keywords") or [])
            values = (
                f"{row_i:02d}",
                game["title"],
                game["release"],
                game["platform"] or "Multi",
                f"{social_bits}\n{seo}".strip(),
            )
            for col, value in enumerate(values):
                fill = "141A2C" if row_i % 2 else CARD
                _shade(games.cell(row_i, col), fill)
                para = games.cell(row_i, col).paragraphs[0]
                para.text = ""
                _run(para, value, size=9 if col == 4 else 10, color=CYAN if col == 4 else INK, bold=col == 1)
                if col == 4:
                    times = game.get("post_times") or {}
                    bits = [
                        times.get("tiktok") and f"TikTok · {times['tiktok']}",
                        times.get("instagram") and f"IG · {times['instagram']}",
                    ]
                    extra = "  |  ".join(bit for bit in bits if bit)
                    if extra:
                        when = games.cell(row_i, col).add_paragraph()
                        _run(when, extra, size=8, color=MUTED)

    note = dark_table(1, 1)
    _shade(note.cell(0, 0), BG)
    _run(
        note.cell(0, 0).paragraphs[0],
        f"Floor Brief · mapped products from the live catalog · {payload['as_of']}",
        size=9,
        color=MUTED,
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    _strip_docx_images(path)


def _strip_docx_images(path: Path) -> None:
    """Drop Word's package thumbnail so the brief ships with no image parts."""
    buf = io.BytesIO()
    with ZipFile(path) as zin, ZipFile(buf, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            name = item.filename
            if name.startswith("docProps/thumbnail") or name.lower().endswith(
                (".jpeg", ".jpg", ".png", ".gif", ".emf", ".wmf")
            ):
                continue
            data = zin.read(name)
            if name == "_rels/.rels":
                data = (
                    data.decode("utf-8")
                    .replace(
                        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/thumbnail" Target="docProps/thumbnail.jpeg"/>',
                        "",
                    )
                    .encode("utf-8")
                )
            elif name == "[Content_Types].xml":
                data = (
                    data.decode("utf-8")
                    .replace('<Default Extension="jpeg" ContentType="image/jpeg"/>', "")
                    .encode("utf-8")
                )
            zout.writestr(item, data)
    path.write_bytes(buf.getvalue())


def configure(*, start: date, end: date, stem: str) -> None:
    global RANGE_START, RANGE_END, STEM
    RANGE_START = start
    RANGE_END = end
    STEM = stem


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="Export a Floor Brief top-event PDF and DOCX")
    parser.add_argument("--start", default=RANGE_START.isoformat())
    parser.add_argument("--end", default=RANGE_END.isoformat())
    parser.add_argument("--stem", default=STEM)
    args = parser.parse_args(argv)
    configure(start=date.fromisoformat(args.start), end=date.fromisoformat(args.end), stem=args.stem)
    payload = build_payload()
    html_path = DOCS / f"{STEM}.html"
    pdf_path = DOCS / f"{STEM}.pdf"
    docx_path = DOCS / f"{STEM}.docx"
    html_path.write_text(render_html(payload), encoding="utf-8")
    write_docx(payload, docx_path)
    write_pdf(payload, pdf_path)
    DESKTOP.mkdir(parents=True, exist_ok=True)
    for src in (pdf_path, docx_path):
        shutil.copy2(src, DESKTOP / src.name)
    print(f"Wrote {pdf_path}")
    print(f"Wrote {docx_path}")
    print(f"Copied to {DESKTOP}")
    for event in payload["events"]:
        games = ", ".join(game["title"] for game in event["games"][:4])
        tags = " ".join(event.get("hashtags") or [])[:80]
        print(f"  {event['name']} [{event['runtime']}] {len(event['games'])} games · {tags} → {games}")


if __name__ == "__main__":
    main()
